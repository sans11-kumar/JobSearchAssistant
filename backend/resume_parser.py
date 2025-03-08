import logging
import os
import PyPDF2
import docx
import re
import nltk
from nltk.corpus import stopwords
import spacy
from nltk.tokenize import word_tokenize

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Download NLTK data
try:
    nltk.download('punkt')
    nltk.download('stopwords')
    nltk.download('averaged_perceptron_tagger')
except Exception as e:
    logger.error(f"Failed to download NLTK data: {e}")
    raise

# Load spaCy model with caching
try:
    nlp = spacy.load('en_core_web_sm')
    logger.info("Loaded spaCy model successfully")
except Exception as e:
    logger.error(f"Failed to load spaCy model: {e}")
    raise

class ResumeParser:
    def __init__(self, resume_path):
        self.resume_path = resume_path
        self.stop_words = set(stopwords.words('english'))
        self.text = self._extract_text()
        
    def _extract_text(self):
        """Extract text from resume file (PDF or DOCX)"""
        file_extension = self.resume_path.split('.')[-1].lower()
        
        if file_extension == 'pdf':
            return self._extract_text_from_pdf()
        elif file_extension in ['docx', 'doc']:
            return self._extract_text_from_docx()
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")
    
    def _extract_text_from_pdf(self):
        """Extract text from PDF file with improved OCR handling"""
        text = ""
        try:
            with open(self.resume_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # First try regular text extraction
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()
                    
                    # If text extraction seems poor (too few words), try OCR fallback
                    if page_text and len(page_text.split()) < 50:
                        try:
                            import pytesseract
                            from pdf2image import convert_from_bytes
                            from PIL import Image
                            
                            # Convert PDF page to image
                            images = convert_from_bytes(file.read(), first_page=page_num+1, last_page=page_num+1)
                            
                            # OCR each image
                            for image in images:
                                page_text = pytesseract.image_to_string(image)
                                text += page_text + "\n"
                        except ImportError:
                            logger.warning("OCR dependencies not installed. Falling back to basic text extraction.")
                            text += page_text + "\n"
                    else:
                        text += page_text + "\n"
            
            # Clean up extracted text
            text = self._clean_extracted_text(text)
            return text
            
        except IOError as e:
            logger.error(f"Error reading PDF file: {e}")
            return ""
        
    def _clean_extracted_text(self, text):
        """Clean and normalize extracted text"""
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Fix common OCR errors
        replacements = {
            r'\b([A-Z])\s([A-Z])\b': r'\1\2',  # Fix spaced capital letters
            r'\b([a-z])\s([a-z])\b': r'\1\2',  # Fix spaced lowercase letters
            r'(\d)\s(\d)': r'\1\2',            # Fix spaced numbers
            r'([a-zA-Z])\s([\.\,\;\:])': r'\1\2',  # Fix spaced punctuation
            r'([\.\,\;\:])\s([a-zA-Z])': r'\1\2'   # Fix spaced punctuation
        }
        
        for pattern, replacement in replacements.items():
            text = re.sub(pattern, replacement, text)
            
        # Normalize line breaks
        text = re.sub(r'[\r\n]+', '\n', text)
        
        return text.strip()
    
    def _extract_text_from_docx(self):
        """Extract text from DOCX file with table and formatting support"""
        text = ""
        try:
            doc = docx.Document(self.resume_path)
            
            # Process paragraphs
            for paragraph in doc.paragraphs:
                # Handle bullet points and numbered lists
                if paragraph.style.name.startswith(('List Paragraph', 'List')):
                    text += "• " + paragraph.text + "\n"
                else:
                    text += paragraph.text + "\n"
            
            # Process tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = self._clean_cell_text(cell.text)
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        text += " | ".join(row_text) + "\n"
                text += "\n"  # Add spacing between tables
            
            # Clean and normalize text
            text = self._clean_extracted_text(text)
            return text
            
        except IOError as e:
            logger.error(f"Error reading DOCX file: {e}")
            return ""
            
    def _clean_cell_text(self, text):
        """Clean and normalize table cell text"""
        # Remove excessive whitespace and line breaks
        text = re.sub(r'\s+', ' ', text)
        text = re.sub(r'[\r\n]+', ' ', text)
        
        # Remove empty cells
        if not text.strip():
            return ""
            
        return text.strip()
    
    def parse(self):
        """Parse resume and extract relevant information"""
        # Extract basic information
        skills = extract_skills(self.text)
        keywords = self.extract_keywords()
        
        # Return structured data
        return {
            'text': self.text,
            'skills': skills,
            'keywords': keywords
        }

    def extract_keywords(self, top_n=30):
        """Extract keywords from resume text using enhanced NLP techniques"""
        # Extract technical terms using custom patterns
        technical_terms = self._extract_technical_terms(self.text)
        
        # Extract RAKE keywords
        rake_keywords = self._extract_rake_keywords(self.text)
        
        # Get noun phrases using spaCy
        doc = nlp(self.text)
        noun_phrases = [chunk.text.lower() for chunk in doc.noun_chunks]
        
        # Combine all terms with weights
        all_terms = []
        
        # Technical terms get highest weight
        for term in technical_terms:
            all_terms.extend([term] * 3)  # Higher weight for technical terms
            
        # RAKE keywords get medium weight
        for term in rake_keywords:
            all_terms.extend([term] * 2)
            
        # Noun phrases get base weight
        all_terms.extend(noun_phrases)
        
        # Calculate term frequencies with weights
        term_freq = {}
        for term in all_terms:
            term_freq[term] = term_freq.get(term, 0) + 1
            
        # Sort by frequency and get top N
        sorted_terms = sorted(term_freq.items(), key=lambda x: x[1], reverse=True)
        
        # Filter and rank keywords
        filtered_terms = []
        for term, freq in sorted_terms:
            if len(term) > 2 and not term.isnumeric():
                # Check if term contains any proper nouns or technical terms
                doc_term = nlp(term)
                if any(token.pos_ in ['PROPN', 'NOUN'] for token in doc_term):
                    # Additional filtering for technical relevance
                    if self._is_technical_term(term):
                        filtered_terms.append(term)
        
        return filtered_terms[:top_n]
        
    def _extract_technical_terms(self, text):
        """Extract technical terms using custom patterns"""
        # Common technical term patterns
        patterns = [
            r'\b[A-Z]{2,}\b',  # Acronyms (AWS, API, etc.)
            r'\b[A-Z][a-z]+[A-Z]\w+\b',  # CamelCase terms
            r'\b\w+\.(js|ts|py|java|cpp|go|rs)\b',  # File extensions
            r'\b\w+-\d+\.\d+\b',  # Version numbers
            r'\b[A-Z][a-z]+(?: [A-Z][a-z]+)*\b'  # Proper nouns
        ]
        
        terms = set()
        for pattern in patterns:
            matches = re.findall(pattern, text)
            terms.update(m.lower() for m in matches)
            
        return list(terms)
        
    def _is_technical_term(self, term):
        """Check if term is likely to be technical"""
        # Common technical suffixes
        technical_suffixes = {
            'ing', 'tion', 'ment', 'ity', 'ance', 'ence', 
            'logy', 'graphy', 'metry', 'scope', 'ware',
            'script', 'kit', 'api', 'sdk', 'ide', 'db'
        }
        
        # Common technical prefixes
        technical_prefixes = {
            'micro', 'macro', 'hyper', 'super', 'multi',
            'inter', 'intra', 'trans', 'auto', 'bio'
        }
        
        # Check suffixes
        for suffix in technical_suffixes:
            if term.endswith(suffix):
                return True
                
        # Check prefixes
        for prefix in technical_prefixes:
            if term.startswith(prefix):
                return True
                
        # Check if term appears in common skills
        if term in extract_skills(self.text):
            return True
            
        return False

    def _extract_rake_keywords(self, text):
        """RAKE implementation for keyword extraction"""
        # Split text into sentences
        sentence_list = nltk.sent_tokenize(text)
        
        # Generate stopwords list
        stop_words = set(stopwords.words('english'))
        stop_words.update(['.', ',', ';', ':', '!', '?', '(', ')', '[', ']', '{', '}'])
        
        # Generate candidate keywords
        phrase_list = []
        for sentence in sentence_list:
            words = nltk.word_tokenize(sentence)
            phrases = self._extract_phrases(words)
            phrase_list.extend(phrases)
        
        # Calculate word scores
        word_freq = {}
        word_degree = {}
        for phrase in phrase_list:
            words = nltk.word_tokenize(phrase)
            phrase_length = len(words)
            for word in words:
                word_freq[word] = word_freq.get(word, 0) + 1
                word_degree[word] = word_degree.get(word, 0) + phrase_length
        
        # Calculate RAKE scores
        word_scores = {}
        for word in word_freq:
            word_scores[word] = word_degree[word] / word_freq[word]
        
        # Calculate phrase scores
        phrase_scores = {}
        for phrase in phrase_list:
            words = nltk.word_tokenize(phrase)
            phrase_score = 0
            for word in words:
                phrase_score += word_scores[word]
            phrase_scores[phrase] = phrase_score
        
        return phrase_scores

def extract_skills(text):
    """Extract skills from resume text using NLP"""
    # Expanded technical skills and keywords
    common_skills = {
        "programming": ["python", "java", "javascript", "c++", "ruby", "php", "swift", "kotlin", "go", "rust", "typescript", "c#", "scala", "r", "matlab"],
        "web": ["html", "css", "react", "angular", "vue", "node.js", "express", "django", "flask", "bootstrap", "sass", "less", "webpack", "graphql", "rest api"],
        "database": ["sql", "mysql", "postgresql", "mongodb", "firebase", "oracle", "nosql", "redis", "cassandra", "neo4j", "dynamodb", "elasticsearch"],
        "cloud": ["aws", "azure", "gcp", "cloud", "docker", "kubernetes", "serverless", "lambda", "ec2", "s3", "cloudformation", "terraform", "ansible"],
        "data": ["machine learning", "data science", "ai", "artificial intelligence", "data analysis", "big data", "tableau", "power bi", "pandas", "numpy", "scikit-learn", "tensorflow", "pytorch", "spark", "hadoop"],
        "tools": ["git", "github", "jira", "agile", "scrum", "ci/cd", "jenkins", "terraform", "ansible", "prometheus", "grafana", "kibana", "splunk"],
        "devops": ["kubernetes", "docker", "helm", "argo cd", "istio", "linkerd", "jenkins", "github actions", "circleci", "gitlab ci"],
        "security": ["owasp", "penetration testing", "vulnerability assessment", "siem", "soc", "firewalls", "vpn", "encryption"],
        "mobile": ["react native", "flutter", "android", "ios", "swift", "kotlin", "xamarin"],
        "testing": ["selenium", "cypress", "jest", "mocha", "junit", "testng", "pytest", "unit testing", "integration testing"]
    }
    
    # Flatten the skills list
    all_skills = [skill for category in common_skills.values() for skill in category]
    
    # Process the text with spaCy
    doc = nlp(text.lower())
    
    # Find skills using regex patterns and spaCy entities
    found_skills = set()
    
    # Check for common skills using regex with fuzzy matching
    for skill in all_skills:
        pattern = r'\b' + re.escape(skill) + r'\b'
        if re.search(pattern, text.lower()):
            found_skills.add(skill)
        else:
            # Try fuzzy matching for slight variations
            pattern = r'\b' + re.escape(skill[:4]) + r'[\w]*\b'
            if re.search(pattern, text.lower()):
                found_skills.add(skill)
    
    # Extract skills from noun phrases with context
    for chunk in doc.noun_chunks:
        chunk_text = chunk.text.lower()
        if chunk_text in all_skills:
            found_skills.add(chunk_text)
        else:
            # Check for multi-word skills
            for skill in all_skills:
                if ' ' in skill and skill in chunk_text:
                    found_skills.add(skill)
    
    # Add named entities that might be technologies
    for ent in doc.ents:
        if ent.label_ in ["PRODUCT", "ORG", "TECH"] and ent.text.lower() in all_skills:
            found_skills.add(ent.text.lower())
    
    # Extract skills from experience sections
    experience_pattern = r'(?:experience|skills|technical skills)[\s\S]*?(?:education|projects|$)'
    experience_match = re.search(experience_pattern, text, re.IGNORECASE)
    if experience_match:
        experience_text = experience_match.group(0)
        for skill in all_skills:
            if skill in experience_text.lower():
                found_skills.add(skill)
    
    # Remove common words that might be false positives
    false_positives = {'data', 'cloud', 'web', 'mobile', 'test'}
    found_skills = {skill for skill in found_skills if skill not in false_positives}
    
    # Sort skills by category for better organization
    categorized_skills = {}
    for category, skills in common_skills.items():
        category_skills = [skill for skill in found_skills if skill in skills]
        if category_skills:
            categorized_skills[category] = category_skills
    
    return categorized_skills
